
# table2_weighted_class_profiling.py

"""
Script to generate Table 2 for manuscript:
Weighted sociodeographic and health characteristics across latent classes (5-class solution).
Outputs a .docx file with weighted percentages (no % symbol) and unweighted counts per class.
"""

import pandas as pd
import numpy as np
from sklearn.preprocessing import StandardScaler
from sklearn.mixture import GaussianMixture
from scipy.stats import chi2_contingency
from docx import Document

# Load dataset
file_path = "gesundheitskompetenz_final_GHP16_PAM_levels_with_GHP16_total.xlsx"
df = pd.read_excel(file_path)

# Define key variables
ghp16_columns = [f"GHP16_{i}_rescal" for i in range(1, 17)]
profile_vars = ['GenderM0F1', 'AgeGroup', 'Education_cat', 'Language',
                'HLS_score_cat', 'PAM_level_cat', 'has_chronic_disease']
class_labels = {
    1: "1. Broadly Moderate",
    2: "2. Globally Low",
    3: "3. Medically Passive",
    4: "4. Peripherally Engaged",
    5: "5. High Engagers"
}

# Add fallback weights if not available
if 'weight' not in df.columns:
    df['weight'] = 1.0

# Prepare GMM for LCA using unweighted data (weights used post hoc only)
X = df[ghp16_columns].dropna()
scaler = StandardScaler()
X_scaled = scaler.fit_transform(X)
gmm = GaussianMixture(n_components=5, covariance_type='full', random_state=42)
gmm.fit(X_scaled)
labels = gmm.predict(X_scaled)
df_profile = df.loc[X.index].copy()
df_profile['LCA_class'] = labels + 1

# Recode language
df_profile["LanguageSimple"] = df_profile["Language"].replace({"D": "German", "I": "Italian"})
df_profile["LanguageSimple"] = df_profile["LanguageSimple"].where(
    df_profile["LanguageSimple"].isin(["German", "Italian"]), "Other/Missing"
)

# Create the Word document
doc = Document()
doc.add_heading("Table 2. Weighted Sociodemographic and Health Characteristics by Latent Class", level=2)
doc.add_paragraph("Values represent weighted percentages (unweighted counts). P-values from chi-square tests on unweighted data. Effect sizes are reported as Cramér's V.")

# Generate table per variable
for var in profile_vars:
    var_to_use = "LanguageSimple" if var == "Language" else var
    count_tab = pd.crosstab(df_profile['LCA_class'], df_profile[var_to_use])
    percent_tab = pd.crosstab(df_profile['LCA_class'], df_profile[var_to_use],
                              values=df_profile.loc[df_profile.index, 'weight'],
                              aggfunc='sum', normalize='index') * 100
    percent_tab = percent_tab.round(1)
    merged_tab = percent_tab.astype(str) + " (" + count_tab.astype(str) + ")"
    merged_tab.index = merged_tab.index.map(class_labels)

    chi2, p, dof, expected = chi2_contingency(count_tab)
    n = count_tab.values.sum()
    phi2 = chi2 / n
    r, k = count_tab.shape
    cramers_v = np.sqrt(phi2 / min(k - 1, r - 1)) if min(k - 1, r - 1) > 0 else np.nan

    doc.add_paragraph(f"{var.replace('_', ' ')} (p = {p:.4f}; Cramér's V = {cramers_v:.3f})", style='Heading 3')
    table = doc.add_table(rows=1, cols=len(merged_tab.columns) + 1)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Class"
    for i, col in enumerate(merged_tab.columns):
        hdr_cells[i + 1].text = str(col)
    for idx, row in merged_tab.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = idx
        for i, val in enumerate(row):
            row_cells[i + 1].text = val

# Save output
output_path = "LCA_Class_Profiling_Weighted_Corrected.docx"
doc.save(output_path)
